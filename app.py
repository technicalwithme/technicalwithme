import streamlit as st
from PIL import Image
from docx import Document
from google import genai
from google.genai import types
import json
import io
import time

# ----------------- ADMIN PASSWORD CONFIGURATION -----------------
ADMIN_PASSWORD = "Sajjad@786"

if "admin_logged_in" not in st.session_state:
    st.session_state.admin_logged_in = False

# Page Configuration
st.set_page_config(
    page_title="Technical With Me | Electrical Engineering Portal",
    page_icon="⚡",
    layout="wide"
)

# Custom Styling (CSS)
st.markdown("""
<style>
    .main-header {
        font-size: 2.2rem;
        font-weight: 700;
        background: linear-gradient(90deg, #1E88E5, #00E676);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0px;
    }
    .sub-text {
        font-size: 1rem;
        color: #6c757d;
        margin-bottom: 20px;
    }
    .ad-card {
        background-color: #f8f9fa;
        border: 2px dashed #00E676;
        border-radius: 10px;
        padding: 15px;
        text-align: center;
        margin: 15px 0;
    }
    .blog-card {
        background: #ffffff;
        border: 1px solid #e0e0e0;
        border-radius: 10px;
        padding: 20px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        margin-bottom: 20px;
    }
</style>
""", unsafe_allow_html=True)

# Session State for Dynamic Blogs
if "blogs" not in st.session_state:
    st.session_state.blogs = [
        {
            "title": "Transformer Tan-Delta & Capacitance Testing Best Practices",
            "category": "Technical",
            "content": "Tan-Delta testing is crucial for assessing insulation degradation in bushings and windings. Maintaining ambient temperature records and using shielded cables prevents noise during 10kV test voltages.",
            "video_url": "https://www.youtube.com/watch?v=dQw4w9WgXcQ"
        }
    ]

# Sidebar Navigation
with st.sidebar:
    st.image("https://img.icons8.com/color/96/lightning-bolt.png", width=60)
    st.title("Technical With Me")
    st.caption("Power Systems & Automation Portal")
    menu = st.radio("Navigation", ["⚡ AI Report Auto-Filler", "📰 Tech Blogs & Vlogs", "🔒 Post New Blog (Admin)", "📢 Sponsor Ads"])
    
    st.divider()
    st.markdown("""
    <div class="ad-card">
        <small style="color: #888;">SPONSORED</small><br>
        <b>Omicron CPC 100 & Testing Kits</b><br>
        <span style="font-size: 12px; color: #555;">Reliable Switchyard Commissioning Tools</span>
    </div>
    """, unsafe_allow_html=True)

# ----------------- PAGE 1: AI AUTO-FILLER -----------------
if menu == "⚡ AI Report Auto-Filler":
    st.markdown('<p class="main-header">⚡ AI Transformer Report Auto-Filler</p>', unsafe_allow_html=True)
    st.markdown('<p class="sub-text">Convert site engineer handwritten sheets directly into structured Word documents.</p>', unsafe_allow_html=True)

    col1, col2 = st.columns([2, 1])

    with col1:
        template_file = st.file_uploader("1. Blank Transformer Format (.docx)", type=["docx"])
        uploaded_report = st.file_uploader("2. Site Engineer Handwritten Sheet (PDF, JPG, PNG)", type=["pdf", "jpg", "png", "jpeg"])

        def get_template_structure(doc):
            structure = []
            for i, p in enumerate(doc.paragraphs):
                txt = p.text.strip()
                if txt:
                    structure.append(f"P[{i}]: {txt}")
            
            for t_idx, tbl in enumerate(doc.tables):
                structure.append(f"\n=== TABLE {t_idx} (Rows: {len(tbl.rows)}, Cols: {len(tbl.columns) if tbl.rows else 0}) ===")
                for r_idx, row in enumerate(tbl.rows):
                    row_content = []
                    for c_idx, cell in enumerate(row.cells):
                        val = cell.text.strip().replace("\n", " ")
                        if not val:
                            row_content.append(f"R{r_idx}C{c_idx}:[EMPTY_WRITEABLE]")
                        else:
                            row_content.append(f"R{r_idx}C{c_idx}:{val}")
                    structure.append(" | ".join(row_content))
            return "\n".join(structure)

        if template_file and uploaded_report:
            if st.button("🚀 Generate Final Report", use_container_width=True):
                api_key = st.secrets.get("GEMINI_API_KEY", "")

                if not api_key:
                    st.error("API Key backend secrets mein nahi mili.")
                else:
                    status = st.empty()
                    status.info("Reading document format and site sheet...")

                    try:
                        doc = Document(template_file)
                        template_map = get_template_structure(doc)

                        if uploaded_report.type == "application/pdf":
                            file_bytes = uploaded_report.getvalue()
                            mime_type = "application/pdf"
                        else:
                            img = Image.open(uploaded_report)
                            if img.mode != 'RGB':
                                img = img.convert('RGB')
                            img.thumbnail((2000, 2000))
                            buf = io.BytesIO()
                            img.save(buf, format='JPEG', quality=90)
                            file_bytes = buf.getvalue()
                            mime_type = "image/jpeg"

                        client = genai.Client(api_key=api_key)
                        file_part = types.Part.from_bytes(data=file_bytes, mime_type=mime_type)

                        status.info("AI extracting all electrical test records...")

                        prompt = f"""
                        You are an expert Transformer Testing & Commissioning Specialist with high-precision OCR abilities.
                        Your task is to thoroughly analyze the handwritten test sheet and map EVERY test reading into the Word document format.

                        ### MANDATORY TRANSFORMER TESTS TO EXTRACT:
                        1. **MAGNETIC BALANCE TEST**:
                           - Voltage applied across phases (e.g. 230V across 1U-1V, 1V-1W, 1W-1U).
                           - Measured induced voltages across other phases (HV and LV windings).
                        2. **MAGNETIZING CURRENT TEST**:
                           - Low-voltage excitation test currents in mA or Amps for all phases (1U, 1V, 1W).
                        3. **SHORT CIRCUIT IMPEDANCE / LOAD TEST**:
                           - Applied Voltage (V), Rated Current (A), Induced Current, % Impedance (%Z), and Loss values.
                        4. **VECTOR GROUP & POLARITY TEST**:
                           - Voltage relationship checks verifying vector configuration (Dyn11, YNd11, etc.) and polarity test results.
                        5. **INSULATION RESISTANCE (IR / MEGGER) & TAN DELTA / CAPACITANCE**:
                           - HV-LV, HV-E, LV-E values (15s, 60s, 600s, PI, DAR).
                           - Bushing C1, C2, and Winding Tan-Delta % and Capacitance (pF).
                        6. **WINDING RESISTANCE**:
                           - Tap-wise resistance across all taps and phase terminals.

                        ### WORD TEMPLATE LAYOUT:
                        {template_map}

                        ### STRICT RULES:
                        - Do NOT truncate output. You MUST capture Magnetic Balance, Magnetizing Current, Polarity, and Short Circuit tables completely.
                        - Only place extracted values into [EMPTY_WRITEABLE] or placeholder cells.
                        - Match test names carefully to Table titles and headers in the template layout.

                        Output strictly a JSON object:
                        {{
                          "paragraph_updates": [
                            {{"index": 0, "append_value": "..."}}
                          ],
                          "table_updates": [
                            {{"table_idx": 0, "row_idx": 1, "col_idx": 2, "value": "..."}}
                          ]
                        }}
                        """

                        candidate_models = ['gemini-3.6-flash']
                        response = None
                        last_error = None

                        for model_name in candidate_models:
                            for attempt in range(3):
                                try:
                                    response = client.models.generate_content(
                                        model=model_name,
                                        contents=[prompt, file_part],
                                        config=types.GenerateContentConfig(
                                            response_mime_type="application/json",
                                            temperature=0.0,
                                            max_output_tokens=8192
                                        )
                                    )
                                    if response and response.text:
                                        break
                                except Exception as err:
                                    last_error = err
                                    err_msg = str(err).lower()
                                    if "503" in err_msg or "unavailable" in err_msg or "resource_exhausted" in err_msg:
                                        time.sleep(4)
                                        continue
                                    else:
                                        raise err
                            if response and response.text:
                                break

                        if not response or not response.text:
                            raise last_error if last_error else Exception("Processing failed. Please retry.")

                        status.info("Writing extracted readings into Word document...")
                        mapping = json.loads(response.text)

                        # Paragraph Updates
                        for p_up in mapping.get("paragraph_updates", []):
                            idx = p_up.get("index")
                            val = p_up.get("append_value", "")
                            if idx is not None and idx < len(doc.paragraphs) and val:
                                doc.paragraphs[idx].text = f"{doc.paragraphs[idx].text} {val}".strip()

                        # Table Updates
                        updated_count = 0
                        for t_up in mapping.get("table_updates", []):
                            t_idx = t_up.get("table_idx")
                            r_idx = t_up.get("row_idx")
                            c_idx = t_up.get("col_idx")
                            val = t_up.get("value", "")

                            if t_idx is not None and t_idx < len(doc.tables):
                                tbl = doc.tables[t_idx]
                                if r_idx is not None and r_idx < len(tbl.rows):
                                    row = tbl.rows[r_idx]
                                    if c_idx is not None and c_idx < len(row.cells):
                                        row.cells[c_idx].text = str(val)
                                        updated_count += 1

                        bio = io.BytesIO()
                        doc.save(bio)
                        status.empty()

                        st.success(f"✅ Report generated successfully! Extracted and filled {updated_count} test parameters.")
                        st.download_button(
                            label="📥 Download Completed Word Document",
                            data=bio.getvalue(),
                            file_name="Completed_Transformer_Report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                    except Exception as e:
                        status.empty()
                        st.error(f"Error: {e}")

    with col2:
        st.markdown("""
        <div class="ad-card">
            <h4>⚡ Industry Solutions</h4>
            <p>High Voltage Transformer Testing & Relay Calibration Services.</p>
            <button style="background-color:#1E88E5; color:white; border:none; padding:8px 16px; border-radius:5px; cursor:pointer;">Contact Experts</button>
        </div>
        """, unsafe_allow_html=True)

# ----------------- PAGE 2: TECH BLOGS & VLOGS -----------------
elif menu == "📰 Tech Blogs & Vlogs":
    st.markdown('<p class="main-header">📰 Technical Articles, Vlogs & News</p>', unsafe_allow_html=True)
    st.write("Latest updates on Power Systems, Substation Automation, Travels, and Industry News.")

    selected_cat = st.radio("Filter By Category:", ["All", "Technical", "Traveling", "News"], horizontal=True)
    filtered_blogs = st.session_state.blogs if selected_cat == "All" else [b for b in st.session_state.blogs if b.get("category") == selected_cat]

    if not filtered_blogs:
        st.info("No posts found in this category.")
    else:
        for blog in filtered_blogs:
            with st.container():
                st.markdown(f"""
                <div class="blog-card">
                    <span style="color:#1E88E5; font-weight:600;">🏷️ {blog.get('category', 'Technical')}</span>
                    <h3 style="margin-top:5px;">{blog['title']}</h3>
                    <p>{blog['content']}</p>
                </div>
                """, unsafe_allow_html=True)
                if blog.get("video_url"):
                    st.video(blog["video_url"])
                st.divider()

# ----------------- PAGE 3: POST NEW BLOG (ADMIN) -----------------
elif menu == "🔒 Post New Blog (Admin)":
    st.markdown('<p class="main-header">✍️ Admin Post Studio</p>', unsafe_allow_html=True)

    if not st.session_state.admin_logged_in:
        st.warning("⚠️ Yeh section password protected hai. Post karne ke liye kripya Admin Password dalein.")
        pwd_input = st.text_input("Enter Admin Password:", type="password")
        
        if st.button("Unlock Admin Panel"):
            if pwd_input == ADMIN_PASSWORD:
                st.session_state.admin_logged_in = True
                st.success("✅ Password correct! Studio unlocked.")
                st.rerun()
            else:
                st.error("❌ Galat password! Kripya sahi password enter karein.")
    else:
        col_admin1, col_admin2 = st.columns([4, 1])
        with col_admin1:
            st.success("🔓 Logged in as Admin")
        with col_admin2:
            if st.button("🚪 Logout"):
                st.session_state.admin_logged_in = False
                st.rerun()

        st.divider()

        with st.form("new_post_form"):
            title = st.text_input("Article / Vlog Title")
            category = st.selectbox("Category", ["Technical", "Traveling", "News"])
            content = st.text_area("Content / Description", height=150)
            video_url = st.text_input("YouTube Video URL (Optional)")
            submitted = st.form_submit_button("📢 Publish Post")
            
            if submitted:
                if title and content:
                    st.session_state.blogs.insert(0, {
                        "title": title,
                        "category": category,
                        "content": content,
                        "video_url": video_url if video_url else None
                    })
                    st.success("🎉 Post published successfully! Check the 'Tech Blogs & Vlogs' section.")
                else:
                    st.warning("Please fill in both title and content.")

# ----------------- PAGE 4: ADS / SPONSORS -----------------
elif menu == "📢 Sponsor Ads":
    st.markdown('<p class="main-header">📢 Sponsorship & Advertising</p>', unsafe_allow_html=True)
    st.write("Monetize your portal using Google AdSense code or direct client banners.")
    
    st.info("💡 To connect Google AdSense: Paste your `<script async src='https://pagead2.googlesyndication.com...'></script>` code here using `st.components.v1.html()`.")
    
    st.markdown("""
    <div class="ad-card" style="padding:40px;">
        <h2>Banner Slot (728x90 / Responsive)</h2>
        <p>Your Google AdSense or Direct Client Banner will appear here.</p>
    </div>
    """, unsafe_allow_html=True)
